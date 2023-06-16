# 这是一个示例 Python 脚本。
from PIL.Image import Resampling
from anki import collection, decks, cards, utils
import zipfile
from docx import Document
import sqlite3
from docx.shared import Pt
import os
from PIL import Image
from docx.shared import Inches

with zipfile.ZipFile("input/input.apkg", 'r') as zip_ref:
    zip_ref.extractall("./temp/")

cwd = "./temp/"
onlyfiles = [os.path.join(cwd, f) for f in os.listdir(cwd) if
             os.path.isfile(os.path.join(cwd, f))]
print(f"process {onlyfiles}")

from docx import Document

document = Document()

break_page = False

file_index = 0

for file in onlyfiles:
    file_index += 1
    is_last = file_index == len(onlyfiles)

    if not os.path.basename(file).isdigit():
        continue
    else:
        print(f"process file {file}")

    try:
        os.rename(file, f"{file}.png")
    except:
        pass

    print(f"get file as:{file}")
    file = f"{file}.png"

    image = Image.open(file)
    image.thumbnail((512, 512), Resampling.HAMMING)
    image.save(file, 'png')

    document.add_picture(file)

    if break_page:
        break_page = False

        # donot add new page at last
        if not is_last:
            document.add_page_break()
    else:
        break_page = True
        document.add_paragraph("\n\n\n\n\n\n\n\n")

# 设置页边距等
document.sections[0].left_margin = (Inches(0.25))
document.sections[0].right_margin = (Inches(0.25))
document.sections[0].top_margin = (Inches(0.25))
document.sections[0].bottom_margin = (Inches(0.25))

# 添加页眉
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

header = document.sections[0].header
paragraph = header.paragraphs[0]

# 设置文字
run = paragraph.add_run()
font = run.font
font.italic = True
font.size = Pt(12)
paragraph.text = "大明酱的Brainstorm.  Powered by MingMoe"
paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

try:
    if os.path.exists("output.docx"):
        os.remove('output.docx')
finally:
    document.save('output.docx')
